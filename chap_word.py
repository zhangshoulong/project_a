#!/usr/bin/env python 
# -*- coding:utf-8 -*-
from docx import Document
from docx.shared import Inches

if __name__ == '__main__':
    document_ins = Document()
    document_ins.add_heading('Document Title', 0)

    document_ins.add_picture('pic_for_word.jpg', width=Inches(6))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document_ins.add_table(rows=1, cols=3, style='Light Grid Accent 2')

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document_ins.save('demo.docx')


