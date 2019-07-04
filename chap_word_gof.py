#!/usr/bin/env python 
# -*- coding:utf-8 -*-
from docx import Document
from docx.shared import Inches, RGBColor
import os, sys

if __name__ == '__main__':
    target_dir = os.path.join(os.path.dirname(sys.argv[0]), 'game_of_thrones')
    with open(os.path.join(target_dir, 'content.txt'), 'r') as file_obj:
        line_list = file_obj.readlines()

        document_ins = Document()
        document_ins.add_heading("For many 'Thrones' fans, season 8 is just the first ending", 0)

        pic_counter = 1
        previous_line = ''
        for each_line in line_list:
            if previous_line == each_line:
                continue

            if '[Photo:'in each_line:
                image_path = os.path.join(target_dir, '%s.jpg' % pic_counter)
                if os.path.exists(image_path):
                    document_ins.add_picture(image_path, width=Inches(6))
                pic_counter = pic_counter + 1
                each_para = document_ins.add_paragraph()
                #step2 italic
                each_run = each_para.add_run(each_line)
                each_run.italic = True
                #step3 set color
                each_run.font.color.rgb = RGBColor(0xAB, 0xAB, 0xAB)

            elif each_line.startswith('\n'):
                continue
            else:
                document_ins.add_paragraph(each_line)

            previous_line = each_line


        document_ins.save('demo_gof.docx')


